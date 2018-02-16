import wx, random, os, time, datetime
from docx import Document

class Example(wx.Frame):
  
    def __init__(self, parent, title):
        super(Example, self).__init__(parent, title=title, 
            size=(460, 200), style= wx.SYSTEM_MENU | wx.CAPTION | wx.CLOSE_BOX)
            
        self.InitUI()
        self.Centre()
        self.Show()     
        
    def InitUI(self):       
    
        panel = wx.Panel(self)

        sizer = wx.GridBagSizer(4, 4)

        hbox = wx.BoxSizer(wx.HORIZONTAL)

        fgs = wx.FlexGridSizer(3, 2, 9, 25)

        title = wx.StaticText(panel, label="ФИО")
        author = wx.StaticText(panel, label="Организация")
        date = wx.StaticText(panel, label="Срок действия (дней)")
        self.tc1 = wx.TextCtrl(panel)
        self.tc2 = wx.TextCtrl(panel)
        self.tc3 = wx.TextCtrl(panel)
                


        fgs.AddMany([(title), (self.tc1, 1, wx.EXPAND), (author), 
            (self.tc2, 1, wx.EXPAND), (date, 1, wx.EXPAND), (self.tc3, 1, wx.EXPAND)])	

        fgs.AddGrowableCol(1, 1)

        buttonGuest = wx.Button(panel, label="GS Гость",pos=(40, 120), size=(90, 28))
        buttonGuestGS = wx.Button(panel, label="Гость",pos=(183, 120), size=(90, 28))
        buttonExit = wx.Button(panel, label="Выход",pos=(320, 120), size=(90, 28))
        self.Bind(wx.EVT_BUTTON, self.guestbutton, buttonGuest)
        self.Bind(wx.EVT_BUTTON, self.guestbuttongs, buttonGuestGS)
        self.Bind(wx.EVT_CLOSE, self.closewindow, buttonExit)
        self.Bind(wx.EVT_BUTTON, self.closebutton, buttonExit)

        hbox.Add(fgs, proportion=1, flag=wx.ALL|wx.EXPAND, border=15)
        panel.SetSizer(hbox)
        

      
    
        
    
    def guestbutton(self, event):
        
        login = []
        chars = 'abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ'
        num = int(5)
        for k in range(1, num+1):
            login.append(random.choice(chars))
        login = "".join(login)

        logvar = self.tc1.GetValue()
        organization = self.tc2.GetValue()
        date = self.tc3.GetValue()
        appenddates = datetime.datetime.now() + datetime.timedelta(days=int(date))
        limiteddate = appenddates.strftime('%d.%m.%Y')
        pw = []
        chars = 'abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ'
        num = int(5)
        for k in range(1, num+1):
            pw.append(random.choice(chars))
        pw = "".join(pw)
    
        file = open('testfilemeow.txt','w+')

        """ Редактирование *.bat файла для группы WiFi-CTS-GS """
        
        file.write('''chcp 1251 >nul
net user %s %s /add /fullname:"%s" /comment:"%s" /expires:"%s"
net localgroup WiFi-CTS-GS %s /add
    ''' % (login, pw, logvar, organization, limiteddate, login)) 
        file.close()
        os.rename('testfilemeow.txt', '%swf.bat' %login)
        opendoc = open('manual.docx', 'rb')
        document = Document(opendoc)
        document.paragraphs[0].insert_paragraph_before('Имя пользователя: «s-sd\%s»            Пароль: «%s»               SSID: CTS-GS' %(login, pw))
        document.save('%s.docx' %login)
        opendoc.close()
        os.startfile('%s.docx' %login, 'print')
        time.sleep(5)
        os.remove('%s.docx' %login)
        wx.MessageBox('Логин: %s                Пароль: %s' %(login, pw), 'Учетная запись создана', wx.OK | wx.ICON_INFORMATION)
        

    def guestbuttongs(self, event):
        login = []
        chars = 'abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ'
        num = int(5)
        for k in range(1, num+1):
            login.append(random.choice(chars))
        login = "".join(login)

        logvar = self.tc1.GetValue()
        organization = self.tc2.GetValue()
        date = self.tc3.GetValue()
        appenddates = datetime.datetime.now() + datetime.timedelta(days=int(date))
        limiteddate = appenddates.strftime('%d.%m.%Y')
        
        pw = []
        chars = 'abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ'
        num = int(5)
        for k in range(1, num+1):
            pw.append(random.choice(chars))
        pw = "".join(pw)
    
        file = open('testfilemeow.txt','w+')

        """ Редактирование *.bat файла для группы WiFi-CTS-Welcome """
        file.write('''chcp 1251 >nul
net user %s %s /add /fullname:"%s" /comment:"%s" /expires:"%s"
net localgroup WiFi-CTS-Welcome %s /add
    ''' % (login, pw, logvar, organization, limiteddate, login)) 
        file.close()
        os.rename('testfilemeow.txt', '%swf.bat' %login)
        opendoc = open('manual.docx', 'rb')
        document = Document(opendoc)
        document.paragraphs[0].insert_paragraph_before('Имя пользователя: «s-sd\%s»            Пароль: «%s»                SSID: CTS-Welcome' %(login, pw))
        document.save('%s.docx' %login)
        opendoc.close()
        os.startfile('%s.docx' %login, 'print')
        time.sleep(5)
        os.remove('%s.docx' %login)
        wx.MessageBox('Логин: %s                Пароль: %s' %(login, pw), 'Учетная запись создана', wx.OK | wx.ICON_INFORMATION)

        

    def closebutton(self, event):
        self.Close(True)

    def closewindow(self, event):
        self.Destroy()


if __name__ == '__main__':
  
    app = wx.App()
    Example(None, title='Создание пользователя WiFi ')
    app.MainLoop()
